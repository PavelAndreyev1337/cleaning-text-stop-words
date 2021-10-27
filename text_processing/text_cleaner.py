from nltk import download
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
from typing import Dict, List, Set, Tuple
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from collections import Counter
from math import sqrt
from itertools import islice


class TextCleaner:
    def __init__(self, language: str = "russian",
                 file_path: str = "text.docx",
                 output_file_path: str = "Андреєв_етап1.docx"):
        download('stopwords')
        download('punkt')
        self.__language = language
        self.__stop_words = set(stopwords.words(self.__language))
        self.__file_path = file_path
        self.__output_file_path = output_file_path
        self.__document = Document(self.__file_path)
        self.__workbook = None
        # used pr-cy.ru {key:word : value:the root of the word, }
        self.__frequent_words = {}
        # the frequency of words in the text {key: word, value: count,}
        self.__text_words_frequency = {}
        # frequency of words in paragraphs {key: word, value: {key:paragraph_number: value: count,},}
        self.__paragraph_words_frequency = {}
        # the number of words in a paragraph {key: paragraph_number, value: count}
        self.__paragraph_words_counts = {}
        # word correlation {key: row_word: {key: column_word,  value: correlation_value,},}
        self.__correlations = {}
        self.__frequent_words_count = 0
        # {key: word, value: {key:paragraph_number: value: count,},} 3-digit accuracy after coma
        self.__relative_frequency = {}
        # {key: kw, value: {key: sw, value: str, key: likelihood_ratio, value: float}, }
        self.__likelihood_ratios = {}
        self.__output_xlsx_file_path = "output.xlsx"
        self.__last_text_words_count = 0
        self.__last_used_stop_words = []

    @property
    def language(self) -> str:
        return self.__language

    @language.setter
    def language(self, language: str) -> None:
        self.__language = language
        self.__stop_words = set(stopwords.words(self.__language))

    @property
    def file_path(self) -> str:
        return self.__file_path

    @property
    def stop_words(self) -> Set[str]:
        return self.__stop_words

    @file_path.setter
    def file_path(self, file_path: str) -> None:
        self.__file_path = file_path
        self.__document = Document(file_path)

    @property
    def last_text_words_count(self) -> int:
        return self.__last_text_words_count

    @property
    def last_used_stop_words(self) -> List[str]:
        return self.__last_used_stop_words

    @property
    def frequent_words(self):
        return self.__frequent_words

    @property
    def text_words_frequency(self):
        return self.__text_words_frequency

    @property
    def paragraph_words_frequency(self):
        return self.__paragraph_words_frequency

    @property
    def paragraph_words_counts(self):
        return self.__paragraph_words_counts

    @property
    def relative_frequency(self):
        return self.__relative_frequency

    @property
    def correlations(self):
        return self.__correlations

    def get_words_count(self):
        return sum([words_count for words_count in self.__paragraph_words_counts.values()])

    def get_last_stop_words_percentage(self) -> float:
        return len(self.__last_used_stop_words) / self.__last_text_words_count

    def __clear_runs(self, paragraph):
        for run in paragraph.runs:
            for word in word_tokenize(run.text):
                self.__last_text_words_count += 1
                if word.lower() in self.__stop_words:
                    self.__last_used_stop_words.append(word)
                    run.text = run.text.replace(f" {word} ", " ")

    def __clear_paragraphs(self):
        for paragraph in self.__document.paragraphs:
            self.__clear_runs(paragraph)

    def __clear_tables(self):
        for table in self.__document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.__clear_runs(paragraph)

    def __reset_results(self):
        self.__last_text_words_count = 0
        self.__last_used_stop_words = []

    def __calculate_words_count(self):
        self.__document = Document(self.__output_file_path)
        for key_word, root in self.__frequent_words.items():
            self.__text_words_frequency[key_word] = 0
            self.__paragraph_words_frequency[key_word] = {}
            paragraph_count = 0
            current_paragraph = ''
            for i, paragraph in enumerate(self.__document.paragraphs):
                current_paragraph += paragraph.text
                words = word_tokenize(current_paragraph)
                if current_paragraph and current_paragraph[0].isdigit() \
                        and ((len(self.__document.paragraphs) - 1 >= i + 1
                              and self.__document.paragraphs[i + 1].text
                              and self.__document.paragraphs[i + 1].text[0].isdigit())
                             or len(self.__document.paragraphs) - 1 == i):
                    self.__paragraph_words_frequency[key_word][paragraph_count] = 0
                    self.__paragraph_words_counts[paragraph_count] = 0
                    self.__paragraph_words_counts[paragraph_count] += len(
                        words)
                    for word in words:
                        if word.startswith(root):
                            self.__text_words_frequency[key_word] += 1
                            self.__paragraph_words_frequency[key_word][paragraph_count] += 1
                    paragraph_count += 1
                    current_paragraph = ''
        self.__workbook.save(self.__output_xlsx_file_path)
        self.__text_words_frequency = {word: count for word, count in
                                       sorted(self.__text_words_frequency.items(),
                                              key=lambda item: item[1],
                                              reverse=True)}

    def __calculate_relative_frequency(self):
        for word, paragraphs in self.__paragraph_words_frequency.items():
            self.__relative_frequency[word] = {}
            for paragraph_number, words_count in paragraphs.items():
                paragraph_words_count = self.__paragraph_words_counts[paragraph_number]
                if not paragraph_words_count:
                    self.__relative_frequency[word][paragraph_number] = 0
                else:
                    self.__relative_frequency[word][paragraph_number] = round(
                        words_count / self.__paragraph_words_counts[paragraph_number], 3)

    def __set_worksheet_cells_width(self, worksheet, size=25):
        for i in range(len(self.__paragraph_words_counts) + 1):  # set width for columns
            worksheet.column_dimensions[get_column_letter(i + 1)].width = size

    def __add_paragraphs_words_frequency_to_xlsx(self,
                                                 worksheet_title: str,
                                                 paragraphs_words_frequency: Dict[str, Dict[int, int]]):
        worksheet = self.__workbook.create_sheet(worksheet_title)
        worksheet.cell(row=1, column=1,
                       value="Кількість слів в абазаці").font = Font(bold=True)
        self.__set_worksheet_cells_width(worksheet)
        for i, words_count in self.__paragraph_words_counts.items():
            worksheet.cell(row=i + 2, column=1, value=words_count)
        for i, frequent_word in enumerate(self.__frequent_words):
            worksheet.cell(row=1, column=i + 2,
                           value=frequent_word).font = Font(bold=True)
            for paragraph_number, words_count in paragraphs_words_frequency[frequent_word].items():
                worksheet.cell(row=paragraph_number + 2,
                               column=i + 2, value=words_count)

    def __add_last_paragraphs_words_frequency_to_xlsx(self,
                                                      frequent_words: Dict[str, str] = {},
                                                      frequent_words_count: int = 21):
        self.__frequent_words_count = frequent_words_count
        if not frequent_words:
            self.__frequent_words = {  # result from pr-cy.ru
                "автосамосвал": "автосамосвал",
                "модель": "модел",
                "транспортный": "транспортн",
                "движение": "движен",
                "карьер": "карьер",

                "работа": "работ",
                "состояние": "состоян",
                "разгрузка": "разгрузк",
                "блок": "блок",
                "пункт": "пункт",

                "погрузка": "погрузк",
                "система": "систем",
                "экскаватор": "экскаватор",
                "время": "врем",
                "управление": "управлен",

                "имитационный": "имитационн",
                "параметр": "параметр",
                "временить": "времен",
                "скорость": "скорост",
                "цикл": "цикл",

                "граф": "граф",
            }
        else:
            self.__frequent_words = frequent_words[:self.__frequent_words_count]
        self.__calculate_words_count()
        self.__calculate_relative_frequency()
        self.__add_paragraphs_words_frequency_to_xlsx(
            "Абсолютна частота", self.__paragraph_words_frequency)
        self.__add_paragraphs_words_frequency_to_xlsx(
            "Відносна частота", self.__relative_frequency)

    def __calculate_relative_frequency_average(self, relative_frequency: Dict[int, int]) -> float:
        return sum(relative_frequency.values()) / len(relative_frequency)

    def __calculate_standard_deviation(self, relative_frequency: Dict[int, int], average: float) -> float:
        return sqrt((1 / (len(relative_frequency) - 1)) *
                    sum([(frequency - average) ** 2 for frequency in relative_frequency.values()]))

    def __add_correlation_to_xlsx(self, worksheet_title: str):
        worksheet = self.__workbook.create_sheet(worksheet_title)
        self.__set_worksheet_cells_width(worksheet, 15)
        for i, row_word in enumerate(self.__frequent_words):
            worksheet.cell(row=1, column=i + 2, value=row_word).font = Font(bold=True)
            for j, column_word in enumerate(self.__frequent_words):
                worksheet.cell(row=j + 2, column=1, value=column_word).font = Font(bold=True)
                if i < j:
                    break
                x_average = self.__calculate_relative_frequency_average(
                    self.__relative_frequency[row_word])
                y_average = self.__calculate_relative_frequency_average(
                    self.__relative_frequency[column_word]
                )
                x_standard_deviation = self.__calculate_standard_deviation(
                    self.__relative_frequency[row_word],
                    x_average)
                y_standard_deviation = self.__calculate_standard_deviation(
                    self.__relative_frequency[column_word],
                    y_average)
                if not row_word in self.__correlations:
                    self.__correlations[row_word] = {}
                correlation = round(sum([
                    ((x_frequency - x_average) / x_standard_deviation)
                    * ((self.__relative_frequency[column_word][paragraph_number] - y_average) / y_standard_deviation)
                    for paragraph_number, x_frequency in self.__relative_frequency[row_word].items()
                ]) * (1 / (len(self.__relative_frequency[row_word]) - 1)), 3)
                worksheet.cell(row=i + 2, column=j + 2, value=correlation)
                self.__correlations[row_word][column_word] = correlation

    def __choice_words_pairs_by_correlation_threshold(self,
                                                      correlation_threshold: float) -> List[Tuple[str, str]]:
        pairs = []  # [(first_word, second_word),]
        for row_word, columns in self.__correlations.items():
            for column_word, correlation in columns.items():
                if row_word != column_word and correlation > correlation_threshold:
                    pairs.append((row_word, column_word))
        return pairs

    def __reduce_words_count_by_matches(self,
                                        pairs: List[Tuple[str, str]],
                                        size: int = 5) -> Dict[str, int]:
        # dictionary of kw {'word': count,}
        return dict(islice(Counter([word for pair in pairs for word in pair]).items(), size))

    def __get_satellite_words(self,
                              keyword: str,
                              lower_correlation_limit: float = 0.15,
                              upper_correlation_limit: float = 0.25) -> str:
        for column_word, correlation in self.__correlations[keyword].items():
            if column_word != keyword \
                    and lower_correlation_limit <= correlation <= upper_correlation_limit:
                    # and correlation > 0.15:
                return column_word

    def __split_paragraphs_to_groups(self,
                                     keyword: str,
                                     satellite_word: str) -> List[List[int]]:
        groups = [[], []]  # the first contains both kw and sw; the second has kw, but no sw
        for paragraph_number, count in self.__paragraph_words_frequency[keyword].items():
            if count:
                satellite_count = self.__paragraph_words_frequency[satellite_word][paragraph_number]
                if satellite_count:
                    groups[0].append(paragraph_number)
                else:
                    groups[1].append(paragraph_number)
        return groups

    def __sum_relative_frequency(self,
                                 keyword: str,
                                 paragraphs: List[int]) -> float:
        return sum([self.__relative_frequency[keyword][paragraph_number] for paragraph_number in paragraphs])

    def __append_likelihood_ratio_to_xlsx(self,
                                          worksheet,
                                          row: int,
                                          keyword: str,
                                          satellite_word: str,
                                          likelihood_ratio: float):
        worksheet.cell(row=row, column=1, value='kw')
        worksheet.cell(row=row, column=2, value=keyword)
        worksheet.cell(row=row + 1, column=1, value='sw')
        worksheet.cell(row=row + 1, column=2, value=satellite_word)
        worksheet.cell(
            row=row + 2,
            column=1,
            value='Відношення правдоподібності').font = Font(bold=True)
        worksheet.cell(
            row=row + 2,
            column=2,
            value=likelihood_ratio).font = Font(bold=True)

    def __add_last_likelihood_ratio_to_xlsx(self,
                                            worksheet_title: str,
                                            correlation_threshold: float = 0.25):
        worksheet = self.__workbook.create_sheet(worksheet_title)
        self.__set_worksheet_cells_width(worksheet, 30)
        pairs = self.__choice_words_pairs_by_correlation_threshold(correlation_threshold)
        keywords = list(self.__reduce_words_count_by_matches(pairs).keys())  # list of kw ['word',]
        row_number = 1
        for i, keyword in enumerate(keywords):
            satellite_word = self.__get_satellite_words(keyword)  # list of sw ['word',]
            if not satellite_word:
                continue
            groups = self.__split_paragraphs_to_groups(keyword, satellite_word)  # [[int, ],[int, ]]
            if not (groups[0] and groups[1]):
                continue
            p1 = self.__sum_relative_frequency(keyword, groups[0])  # P(kwi/swj)
            p2 = self.__sum_relative_frequency(keyword, groups[1])  # P(kwi/¬swj)
            likelihood_ratio = p1 / p2
            self.__likelihood_ratios[keyword] = {
                'sw': satellite_word,
                'likelihood_ratio': likelihood_ratio
            }
            self.__append_likelihood_ratio_to_xlsx(
                worksheet,
                row_number,
                keyword,
                satellite_word,
                likelihood_ratio
            )
            row_number += 3

    def export_to_xlsx(self,
                       frequent_words: Dict[str, str] = {},
                       output_xlsx_file_path: str = "Андреєв_етап3.xlsx",
                       frequent_words_count: int = 21):  # stage 3
        self.__output_xlsx_file_path = output_xlsx_file_path
        self.__workbook = Workbook()
        self.__add_last_paragraphs_words_frequency_to_xlsx(frequent_words, frequent_words_count)
        self.__add_correlation_to_xlsx("Кореляція")
        self.__add_last_likelihood_ratio_to_xlsx('Відношення правдоподібності')
        self.__workbook.remove(self.__workbook[self.__workbook.sheetnames[0]])
        self.__workbook.save(self.__output_xlsx_file_path)

    def clear(self):
        self.__reset_results()
        self.__clear_paragraphs()
        self.__clear_tables()
        self.__document.save(self.__output_file_path)
